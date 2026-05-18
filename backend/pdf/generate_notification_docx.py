import sys
import io
import os
import json
import tempfile
import requests

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.abspath(os.path.join(BASE_DIR, ".."))

MODELO_PADRAO = os.path.join(BASE_DIR, "modelo-notificacao-dinamico.docx")
SAIDA_PADRAO = os.path.join(BASE_DIR, "notificacao_preenchida.docx")


def to_s(v):
    return "" if v is None else str(v)


def as_float(x, default=0.0):
    try:
        if isinstance(x, str):
            x = x.replace(",", ".").strip()
        return float(x)
    except Exception:
        return default


def fmt2(x):
    return f"{as_float(x, 0):.2f}"


def clamp_nota(n):
    n = as_float(n, 0)
    return max(0.0, min(10.0, round(n, 2)))


def classificar(_nota, dados):
    return to_s(
        dados.get("classificacaoComportamental")
        or dados.get("comportamento")
    ) or "—"


def aplicar_formatacao_paragrafo(paragraph):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def substituir_tokens_no_texto(texto, replacements):
    for key, value in replacements.items():
        texto = texto.replace("{{" + key + "}}", to_s(value))
    return texto


def substituir_no_paragrafo(paragraph, replacements):
    full_text = paragraph.text or ""
    novo_texto = substituir_tokens_no_texto(full_text, replacements)

    if novo_texto == full_text:
        aplicar_formatacao_paragrafo(paragraph)
        return

    paragraph.clear()
    partes = novo_texto.split("\n")

    for i, parte in enumerate(partes):
        if i > 0:
            paragraph.add_run().add_break()

        run = paragraph.add_run(parte)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def substituir_em_tabela(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                substituir_no_paragrafo(p, replacements)
                aplicar_formatacao_paragrafo(p)

            for inner_table in cell.tables:
                substituir_em_tabela(inner_table, replacements)


def substituir_em_documento(doc, replacements):
    for p in doc.paragraphs:
        substituir_no_paragrafo(p, replacements)
        aplicar_formatacao_paragrafo(p)

    for table in doc.tables:
        substituir_em_tabela(table, replacements)


def limpar_paragrafos(paragraphs):
    for p in paragraphs:
        p.clear()


def remover_bordas_tabela(table):
    try:
        tbl = table._tbl
        tbl_pr = tbl.tblPr

        for child in list(tbl_pr):
            if child.tag.endswith("tblBorders"):
                tbl_pr.remove(child)
    except Exception:
        pass


def resolver_imagem(url_ou_caminho):
    valor = to_s(url_ou_caminho).strip()
    if not valor:
        return None

    try:
        if valor.startswith("http://") or valor.startswith("https://"):
            response = requests.get(valor, timeout=10)
            response.raise_for_status()

            suffix = ".png"
            lower = valor.lower()

            if ".jpg" in lower or ".jpeg" in lower:
                suffix = ".jpg"
            elif ".webp" in lower:
                suffix = ".webp"

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(response.content)
            tmp.close()

            return tmp.name

        candidatos = []

        if valor.startswith("/uploads/"):
            relativo = valor.replace("/uploads/", "")
            candidatos.append(os.path.join(ROOT_DIR, "uploads", relativo))
            candidatos.append(os.path.join(ROOT_DIR, "public", "uploads", relativo))

        if valor.startswith("uploads/"):
            relativo = valor.replace("uploads/", "")
            candidatos.append(os.path.join(ROOT_DIR, "uploads", relativo))
            candidatos.append(os.path.join(ROOT_DIR, "public", "uploads", relativo))

        candidatos.append(os.path.join(ROOT_DIR, valor.lstrip("/")))
        candidatos.append(valor)

        for caminho in candidatos:
            if caminho and os.path.exists(caminho):
                return caminho

    except Exception as e:
        print(f"[DOCX][IMAGEM] Falha ao resolver imagem: {e}", file=sys.stderr)

    return None


def inserir_imagem_na_celula(cell, caminho, largura_cm=1.8):
    if not caminho or not os.path.exists(caminho):
        return

    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        run = p.add_run()
        run.add_picture(caminho, width=Cm(largura_cm))

    except Exception as e:
        print(f"[DOCX][IMAGEM] Falha ao inserir brasão: {e}", file=sys.stderr)


def montar_cabecalho_word(doc, dados):
    orgao = to_s(dados.get("orgaoSuperior")).upper().strip()
    instituicao = to_s(dados.get("nomeInstituicao")).upper().strip()
    subtitulo = to_s(dados.get("subtituloInstitucional")).strip()

    mostrar_esq = dados.get("mostrarBrasaoEsquerdo", True) is not False
    mostrar_dir = dados.get("mostrarBrasaoDireito", True) is not False

    brasao_esq = resolver_imagem(dados.get("brasaoEsquerdoUrl")) if mostrar_esq else None
    brasao_dir = resolver_imagem(dados.get("brasaoDireitoUrl")) if mostrar_dir else None

    for section in doc.sections:
        section.different_first_page_header_footer = False

        # Cabeçalho mais compacto e mais próximo do topo.
        section.header_distance = Cm(0.1)
        section.top_margin = Cm(1.55)

        header = section.header
        header.is_linked_to_previous = False

        limpar_paragrafos(header.paragraphs)

        # Tabela mais larga, com brasões mais laterais e texto central mais espalhado.
        table = header.add_table(rows=1, cols=3, width=Cm(18.5))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        table.columns[0].width = Cm(4.2)
        table.columns[1].width = Cm(10.1)
        table.columns[2].width = Cm(4.2)

        remover_bordas_tabela(table)

        left = table.cell(0, 0)
        center = table.cell(0, 1)
        right = table.cell(0, 2)

        inserir_imagem_na_celula(left, brasao_esq, 1.8)
        inserir_imagem_na_celula(right, brasao_dir, 1.8)

        center.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = center.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        def add_line(texto, tamanho=10, bold=False):
            run = p.add_run(texto)
            run.bold = bold
            run.font.name = "Times New Roman"
            run.font.size = Pt(tamanho)
            return run

        if orgao:
            add_line(orgao, tamanho=9, bold=True)

        if instituicao:
            if orgao:
                p.add_run().add_break()
            add_line(instituicao, tamanho=11, bold=True)

        if subtitulo:
            p.add_run().add_break()
            add_line(subtitulo, tamanho=9, bold=False)


def montar_rodape_word(doc, dados):
    if dados.get("mostrarRodape") is False:
        return

    texto = to_s(dados.get("rodapeInstitucional")).strip()
    if not texto:
        return

    for section in doc.sections:
        section.footer_distance = Cm(0.6)
        section.bottom_margin = Cm(1.8)

        footer = section.footer
        footer.is_linked_to_previous = False

        limpar_paragrafos(footer.paragraphs)

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        run = p.add_run(texto)
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)


def main():
    try:
        raw = sys.stdin.read()
        dados = json.loads(raw) if raw else {}
    except Exception as e:
        print(f"Erro ao ler JSON: {e}", file=sys.stderr)
        sys.exit(1)

    modelo_path = dados.get("modeloPath") or MODELO_PADRAO
    saida_path = dados.get("saidaPath") or SAIDA_PADRAO

    if not os.path.isabs(modelo_path):
        modelo_path = os.path.join(BASE_DIR, modelo_path)

    if not os.path.isabs(saida_path):
        saida_path = os.path.join(BASE_DIR, saida_path)

    if not os.path.exists(modelo_path):
        print(f"Modelo não encontrado: {modelo_path}", file=sys.stderr)
        sys.exit(1)

    aluno_nome = to_s(dados.get("alunoNome") or dados.get("aluno"))
    turma = to_s(dados.get("turma") or dados.get("alunoTurma"))
    numero = to_s(dados.get("numeroSequencial") or dados.get("numero"))
    data = to_s(dados.get("dataPorExtenso") or dados.get("dataHora"))
    descricao = to_s(dados.get("descricaoInfracao") or "—")
    observacao = to_s(dados.get("observacao") or "-")

    nota_anterior = clamp_nota(dados.get("notaAnterior"))
    nota_atual = clamp_nota(dados.get("notaAtual"))
    classificacao = classificar(nota_atual, dados)

    delta_raw = dados.get("valorNumerico", None)

    if delta_raw is None:
        delta = round(nota_atual - nota_anterior, 2)
    else:
        delta = round(as_float(delta_raw, 0), 2)

    natureza = to_s(dados.get("natureza")).strip().lower()
    if not natureza:
        natureza = "elogio" if delta > 0 else "indisciplina"

    texto_institucional = to_s(dados.get("textoInstitucional"))
    nome_regulamento = to_s(dados.get("regulamentoNome"))

    cidade = to_s(dados.get("cidade"))
    estado = to_s(dados.get("estado"))

    if natureza == "elogio":
        titulo = "ELOGIO INDIVIDUAL"
        frase_resultado = (
            f"Este reconhecimento resultou em acréscimo de {abs(delta):.2f} pontos, "
            f"enquadrando o(a) aluno(a) no comportamento {classificacao}."
        )
        frase_final = (
            "Parabenizamos pela postura e incentivamos a continuidade desse desempenho."
        )
    else:
        titulo = "NOTIFICAÇÃO DISCIPLINAR"
        frase_resultado = (
            f"Esta ocorrência resultou em redução de {abs(delta):.2f} pontos, "
            f"enquadrando o(a) aluno(a) no comportamento {classificacao}."
        )
        frase_final = (
            "Reforçamos a importância do cumprimento das normas institucionais."
        )

    replacements = {
        "cabecalho": "",
        "tituloDocumento": titulo,
        "titulo": titulo,
        "regulamentoNome": nome_regulamento,
        "textoInstitucional": texto_institucional,
        "numeroSequencial": numero,
        "numero": numero,
        "aluno": aluno_nome,
        "alunoNome": aluno_nome,
        "turma": turma,
        "dataHora": data,
        "dataPorExtenso": data,
        "descricaoInfracao": descricao,
        "observacao": observacao,
        "notaAnterior": fmt2(nota_anterior),
        "notaAtual": fmt2(nota_atual),
        "notaFinal": fmt2(nota_atual),
        "comportamento": classificacao,
        "fraseResultado": frase_resultado,
        "fraseFinal": frase_final,
        "cidade": cidade,
        "estado": estado,
    }

    try:
        doc = Document(modelo_path)

        montar_cabecalho_word(doc, dados)
        montar_rodape_word(doc, dados)

        substituir_em_documento(doc, replacements)

        doc.save(saida_path)

    except Exception as e:
        print(f"Erro ao gerar DOCX: {e}", file=sys.stderr)
        sys.exit(1)

    print(saida_path)


if __name__ == "__main__":
    main()