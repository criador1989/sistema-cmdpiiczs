import sys
import io
import os
import json
import hashlib
import requests

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DEFAULT_TEMPLATE = os.path.join(BASE_DIR, "templates", "termo_ciencia.docx")
DEFAULT_OUTPUT = os.path.join(BASE_DIR, "output", "documento_gerado.docx")


def to_s(v):
    return "" if v is None else str(v)


def now_br():
    return datetime.now().strftime("%d/%m/%Y %H:%M")


def gerar_hash(dados):
    try:
        raw = json.dumps(dados, ensure_ascii=False, sort_keys=True)
        return hashlib.sha256(raw.encode("utf-8")).hexdigest()
    except Exception:
        return ""


def aplicar_formatacao_paragrafo(paragraph):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def substituir_tokens(texto, replacements):
    for key, value in replacements.items():
        texto = texto.replace("{{" + key + "}}", to_s(value))
    return texto


def substituir_paragrafo(paragraph, replacements):
    texto_original = paragraph.text or ""
    texto_novo = substituir_tokens(texto_original, replacements)

    if texto_original == texto_novo:
        aplicar_formatacao_paragrafo(paragraph)
        return

    paragraph.clear()
    run = paragraph.add_run(texto_novo)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def substituir_tabela(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                substituir_paragrafo(p, replacements)
                aplicar_formatacao_paragrafo(p)

            for inner in cell.tables:
                substituir_tabela(inner, replacements)


def substituir_documento(doc, replacements):
    for p in doc.paragraphs:
        substituir_paragrafo(p, replacements)
        aplicar_formatacao_paragrafo(p)

    for table in doc.tables:
        substituir_tabela(table, replacements)

    for section in doc.sections:
        for p in section.header.paragraphs:
            substituir_paragrafo(p, replacements)

        for table in section.header.tables:
            substituir_tabela(table, replacements)

        for p in section.footer.paragraphs:
            substituir_paragrafo(p, replacements)

        for table in section.footer.tables:
            substituir_tabela(table, replacements)


def baixar_imagem(url):
    if not url:
        return None

    try:
        limpa = str(url).strip()

        if limpa.startswith("/uploads/"):
            caminho = os.path.abspath(os.path.join(BASE_DIR, "..", limpa.lstrip("/")))
            if os.path.exists(caminho):
                return open(caminho, "rb")
            print(f"Imagem local não encontrada: {caminho}", file=sys.stderr)
            return None

        if limpa.startswith("uploads/"):
            caminho = os.path.abspath(os.path.join(BASE_DIR, "..", limpa))
            if os.path.exists(caminho):
                return open(caminho, "rb")
            print(f"Imagem local não encontrada: {caminho}", file=sys.stderr)
            return None

        response = requests.get(limpa, timeout=10)
        response.raise_for_status()
        return BytesIO(response.content)

    except Exception as e:
        print(f"Erro ao carregar imagem institucional: {e}", file=sys.stderr)
        return None


def limpar_paragrafos(container):
    for p in container.paragraphs:
        p.clear()


def add_linha_cabecalho(paragraph, texto, tamanho=9, negrito=True):
    if not texto:
        return

    r = paragraph.add_run(texto.upper())
    r.bold = negrito
    r.font.name = "Times New Roman"
    r.font.size = Pt(tamanho)
    paragraph.add_run("\n")


def aplicar_cabecalho_institucional(doc, dados):
    identidade = dados.get("identidadeInstitucional") or {}

    orgao = to_s(dados.get("orgaoSuperior") or identidade.get("orgaoSuperior"))
    nome = to_s(dados.get("instituicaoNome") or identidade.get("nomeInstituicao") or dados.get("cabecalho"))

    subtitulo = to_s(
        dados.get("subtituloInstitucional") or
        identidade.get("subtituloInstitucional") or
        identidade.get("subtitulo")
    )

    cidade = to_s(dados.get("cidade") or identidade.get("cidade"))
    uf = to_s(dados.get("estado") or dados.get("uf") or identidade.get("uf"))

    if nome and cidade and uf and cidade.upper() not in nome.upper():
        nome = f"{nome} - {cidade}/{uf}"

    brasao_esq_url = dados.get("brasaoEsquerdoUrl") or identidade.get("brasaoEsquerdoUrl")
    brasao_dir_url = dados.get("brasaoDireitoUrl") or identidade.get("brasaoDireitoUrl")

    mostrar_esq = dados.get("mostrarBrasaoEsquerdo", identidade.get("mostrarBrasaoEsquerdo", True))
    mostrar_dir = dados.get("mostrarBrasaoDireito", identidade.get("mostrarBrasaoDireito", True))

    for section in doc.sections:
        section.top_margin = Inches(0.42)
        section.header_distance = Inches(0.10)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        header = section.header
        limpar_paragrafos(header)

        table = header.add_table(rows=1, cols=3, width=Inches(7.0))
        table.autofit = False
        for row in table.rows:
            row.height = Inches(0.95)

        table.columns[0].width = Inches(1.15)
        table.columns[1].width = Inches(4.70)
        table.columns[2].width = Inches(1.15)

        left = table.cell(0, 0)
        center = table.cell(0, 1)
        right = table.cell(0, 2)

        for cell in [left, center, right]:
            cell.width = Inches(1.15)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 0.9

        center.width = Inches(4.70)

        if mostrar_esq and brasao_esq_url:
            img = baixar_imagem(brasao_esq_url)
            if img:
                p = left.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(0)

                run = p.add_run()
                run.add_picture(img, width=Inches(0.54))

        p = center.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 0.9

        add_linha_cabecalho(p, orgao, tamanho=8, negrito=True)
        add_linha_cabecalho(p, nome, tamanho=8, negrito=True)

        if subtitulo:
            r = p.add_run(subtitulo)
            r.font.name = "Times New Roman"
            r.font.size = Pt(8)

        if mostrar_dir and brasao_dir_url:
            img = baixar_imagem(brasao_dir_url)
            if img:
                p = right.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(0)

                run = p.add_run()
                run.add_picture(img, width=Inches(0.54))


def aplicar_rodape_institucional(doc, dados):
    identidade = dados.get("identidadeInstitucional") or {}

    mostrar = dados.get("mostrarRodape", identidade.get("mostrarRodape", True))
    if mostrar is False:
        return

    texto = to_s(
        dados.get("rodapeInstitucional") or
        identidade.get("rodapeInstitucional") or
        identidade.get("rodapePadrao") or
        dados.get("textoRodape")
    ).strip()

    if not texto:
        return

    for section in doc.sections:
        footer = section.footer
        limpar_paragrafos(footer)

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        r = p.add_run(texto)
        r.font.name = "Times New Roman"
        r.font.size = Pt(8)


def main():
    try:
        raw = sys.stdin.read()
        dados = json.loads(raw) if raw else {}
    except Exception as e:
        print(f"Erro JSON: {e}", file=sys.stderr)
        sys.exit(1)

    template_path = dados.get("templatePath") or DEFAULT_TEMPLATE
    output_path = dados.get("outputPath") or DEFAULT_OUTPUT

    if not os.path.isabs(template_path):
        template_path = os.path.join(BASE_DIR, template_path)

    if not os.path.isabs(output_path):
        output_path = os.path.join(BASE_DIR, output_path)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if not os.path.exists(template_path):
        print(f"Template não encontrado: {template_path}", file=sys.stderr)
        sys.exit(1)

    hash_documento = gerar_hash(dados)

    replacements = {
        "instituicaoNome": to_s(dados.get("instituicaoNome")),
        "cabecalho": to_s(dados.get("cabecalho")),
        "cidade": to_s(dados.get("cidade")),
        "estado": to_s(dados.get("estado")),
        "numeroProcesso": to_s(dados.get("numeroProcesso")),
        "alunoNome": to_s(dados.get("alunoNome")),
        "turma": to_s(dados.get("turma")),
        "responsavelNome": to_s(dados.get("responsavelNome")),
        "responsavelParentesco": to_s(dados.get("responsavelParentesco")),
        "descricaoFato": to_s(dados.get("descricaoFato")),
        "providencias": to_s(dados.get("providencias")),
        "dataFato": to_s(dados.get("dataFato")),
        "horaFato": to_s(dados.get("horaFato")),
        "localFato": to_s(dados.get("localFato")),
        "naturezaProcedimento": to_s(dados.get("naturezaProcedimento")),
        "classificacaoOcorrencia": to_s(dados.get("classificacaoOcorrencia")),
        "gravidade": to_s(dados.get("gravidade")),
        "statusProcesso": to_s(dados.get("statusProcesso")),
        "dataCiencia": to_s(dados.get("dataCiencia")),
        "ipCiencia": to_s(dados.get("ipCiencia")),
        "respostaResponsavel": to_s(dados.get("respostaResponsavel")),
        "resultadoAcompanhamento": to_s(dados.get("resultadoAcompanhamento")),
        "parecerFinal": to_s(dados.get("parecerFinal")),
        "dataGeracao": now_br(),
        "hashDocumento": hash_documento,
        "usuarioGerador": to_s(dados.get("usuarioGerador")),
        "assinaturaDigital": f"Documento gerado digitalmente pelo Axoriin em {now_br()}",
        "textoRodape": (
            dados.get("rodapeInstitucional") or
            "Documento institucional gerado eletronicamente pela plataforma Axoriin."
        )
    }

    try:
        doc = Document(template_path)

        aplicar_cabecalho_institucional(doc, dados)
        aplicar_rodape_institucional(doc, dados)

        substituir_documento(doc, replacements)

        doc.save(output_path)

    except Exception as e:
        print(f"Erro ao gerar documento: {e}", file=sys.stderr)
        sys.exit(1)

    print(output_path)


if __name__ == "__main__":
    main()